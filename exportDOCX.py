data1 = [
    "List of Topics",
    [
        ["Topic1", [
            ["descriptorAlpha", "re:^*", [
                ["Subtopic1A", [
                    ["descriptorBeta", "re:^*", [
                        ["SubSubtopic1A1"],
                        ["SubSubtopic1A2"],
                        ["SubSubtopic1A3"]
                    ]]
                ]],
                ["Subtopic1B", [
                    ["descriptorGamma", "re:^*", [
                        ["SubSubtopic1B1"],
                        ["SubSubtopic1B2"],
                        ["SubSubtopic1B3"]
                    ]]
                ]]
            ]]
        ]],
        ["Topic2", [
            ["descriptorDelta", "re:^*", [
                ["Subtopic2A", [
                    ["descriptorEpsilon", "re:^*", [
                        ["SubSubtopic2A1"],
                        ["SubSubtopic2A2"],
                        ["SubSubtopic2A3"]
                    ]]
                ]],
                ["Subtopic2B", [
                    ["descriptorZeta", "re:^*", [
                        ["SubSubtopic2B1"],
                        ["SubSubtopic2B2"],
                        ["SubSubtopic2B3"]
                    ]]
                ]]
            ]]
        ]]
    ],
    [
        ["Source1", "web"],
        ["Source2", "book"],
        ["Source3", "journal"],
        ["Source4", "publication"]
    ]
]
data2 = [
    "Parathyroid Physiology", 
    [
        ["Parathyroid Hormone", [
            ["Function", "of^*", [
                ["Calcium Regulation", [
                    ["Effects", "via^*", [
                        ["Increase serum calcium"]
                    ]]
                ]]
            ]],
            ["Synthesis", "of^*", [
                ["Gene Transcription", [
                    ["Control", "by^*", [
                        ["Hypocalcemia stimulates via CaSR"],
                        ["Eucalcemia inhibits via CaSR"],
                        ["Vitamin D inhibits via VDR"]
                    ]]
                ]]
            ]],
            ["Secretion", "of^*", [
                ["CaSR Signaling", [
                    ["Regulation", "through^*", [
                        ["Hypocalcemia stimulates"],
                        ["Hypercalcemia inhibits"],
                        ["Hyperphosphatemia stimulates"]
                    ]]
                ]],
                ["Magnesium Effects", [
                    ["Response", "to^*", [
                        ["Mild deficiency stimulates"],
                        ["Severe deficiency inhibits"]
                    ]]
                ]]
            ]],
            ["Receptors", "of^*", [
                ["PTH1R", [
                    ["Location", "in^*", [
                        ["Bone", [
                            ["Actions", "via^*", [
                                ["Osteoblast activation"],
                                ["Osteoclast stimulation"],
                                ["Calcium liberation"]
                            ]]
                        ]],
                        ["Kidney", [
                            ["Actions", "via^*", [
                                ["Calcium reabsorption"],
                                ["Phosphate excretion"],
                                ["Vitamin D activation"]
                            ]]
                        ]]
                    ]]
                ]]
            ]]
        ]]
    ],
    [
        ["Regulation", "of^*"],
        ["Target", "of^*"],
        ["Effect", "of^*"]
    ]
]

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from mapCrude import create_map, get_content_from_path, reorder_map

def refine_map_for_docx(data):
    map_crude_reordered = reorder_map(create_map(data))
    map_refined = []
    
    for path in map_crude_reordered:
        content = get_content_from_path(data, path)
        if path and path[0] != 2:  # Exclude bibliography section
            if isinstance(content, str) and len(path) < 4:  # Append paths for metatopic and topics
                map_refined.append((path, content, "topics"))
            elif isinstance(content, str):
                sibling_content = get_content_from_path(data, path[:-1] + [1])
                if "^*" in sibling_content and "^*" not in content:  # Append paths for descriptors with a sibling containing "^*" !!!
                    map_refined.append((path, content, "descriptors"))
                parent_content = get_content_from_path(data, path[:-3] + [1])
                if isinstance(parent_content, str) and "^*" in parent_content.lower():  # Append paths for subtopics with a parent containing "^*" !!!
                    map_refined.append((path, content, "subtopics"))
    return map_refined

def export_to_docx(data):
    # Create a new document
    document = Document()

    # Get the refined map for the document
    map_refined = refine_map_for_docx(data)

    # Define a function to set paragraph indentation
    def set_paragraph_indentation(paragraph, level):
        p = paragraph._element
        pPr = p.get_or_add_pPr()
        indentation = OxmlElement('w:ind')
        indentation.set(qn('w:left'), str(level * 180))  # 720 twips = 0.5 inch per level
        pPr.append(indentation)

    for path, content, content_type in map_refined:
        if content_type == "topics":
            # Add a heading
            heading_level = len(path) - 1
            document.add_heading(content, level=heading_level)
        elif content_type == "descriptors":
            # Add a paragraph
            paragraph = document.add_paragraph(content, style='List Bullet')
            # Set indentation based on path length
            set_paragraph_indentation(paragraph, len(path) - 1)
            # Italicize the paragraph if it is a descriptor
            if "^*" in content:
                run = paragraph.runs[0]
                run.italic = True
        elif content_type == "subtopics":
            # Add a paragraph
            paragraph = document.add_paragraph(content, style='List Bullet')
            # Set indentation based on path length
            set_paragraph_indentation(paragraph, len(path) - 1)
            # Italicize the paragraph if it is a descriptor
            if "^*" in content:
                run = paragraph.runs[0]
                run.italic = True

    # Save the document with the name from data[0]
    document_name = f"{data[0]}.docx"
    document.save(document_name)
    print(f"Document saved as '{document_name}'")

export_to_docx(data2)